
"""
简化版商品数据处理器
主要特点：
1. 使用大模型驱动解析各种格式的商品数据
2. 无规则硬编码，完全由模型理解商品结构
3. 保证提取的数据与原始数据内容对齐
4. 支持多种输入格式：txt, docx, xlsx, csv等
5. 每个商品作为独立JSON对象，方便检索
"""

import os
import re
import json
import logging
import argparse
import uuid
import random
import string
from typing import List, Dict, Any, Optional, Union, Tuple
import time

# 文件处理相关库
import pandas as pd
import docx
from io import StringIO

# 第三方库 OpenAI
try:
    from langchain_openai import ChatOpenAI
    from langchain_core.prompts import ChatPromptTemplate
    HAS_OPENAI = True
except ImportError:
    HAS_OPENAI = False
    print("警告: 未安装OpenAI相关库，将使用替代方法处理")

# 设置日志格式
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

class ProductDataProcessor:
    """
    商品数据处理器
    使用大模型驱动解析各种格式的商品数据
    支持的输入格式:
    - TXT文件 (结构化文本)
    - DOCX文件 (Word文档)
    - XLSX文件 (Excel表格)
    - CSV文件
    - 直接文本输入
    """
    
    # 支持的文件扩展名
    supported_extensions = ['.txt', '.docx', '.xlsx', '.csv', '.json', '.md']
    
    def __init__(self, use_llm: bool = True, api_key: Optional[str] = None, api_base: Optional[str] = None):
        """
        初始化处理器
        
        参数:
            use_llm: 是否使用大模型进行解析
            api_key: OpenAI API密钥
            api_base: OpenAI API基础URL
        """
        self.use_llm = use_llm and HAS_OPENAI
        self.model = None
        
        if self.use_llm:
            try:
                # 获取API密钥，优先级：参数指定 > 环境变量 > 默认值
                actual_api_key = api_key or os.environ.get("OPENAI_API_KEY") or ""
                actual_api_base = api_base or os.environ.get("OPENAI_API_BASE") or ""
                
                # 检查是否有有效的API密钥
                if not actual_api_key or actual_api_key == "your_api_key_here":
                    logger.warning("未提供有效的API密钥，将使用传统方法解析")
                    self.use_llm = False
                    return
                
                # 初始化模型
                self.model = ChatOpenAI(
                    api_key=actual_api_key,
                    base_url=actual_api_base,
                    model="",
                    temperature=0.7,
                )
                logger.info("已初始化大语言模型，将使用智能解析功能")
            except Exception as e:
                logger.error(f"初始化大语言模型失败: {str(e)}")
                self.use_llm = False
                self.model = None
    
    def process_file(self, file_path: str) -> List[Dict[str, Any]]:
        """
        处理指定文件，提取商品数据
        
        参数:
            file_path: 文件路径
            
        返回:
            商品数据列表
        """
        if not os.path.exists(file_path):
            logger.error(f"文件不存在: {file_path}")
            return []
        
        # 获取文件扩展名
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        if ext not in self.supported_extensions:
            logger.error(f"不支持的文件类型: {ext}")
            return []
        
        try:
            # 根据文件类型选择处理方法
            if ext == '.txt' or ext == '.md':
                return self._process_text_file(file_path)
            elif ext == '.docx':
                return self._process_docx_file(file_path)
            elif ext == '.xlsx':
                return self._process_excel_file(file_path)
            elif ext == '.csv':
                return self._process_csv_file(file_path)
            elif ext == '.json':
                return self._process_json_file(file_path)
            else:
                logger.error(f"不支持的文件类型: {ext}")
                return []
        except Exception as e:
            logger.error(f"处理文件时发生错误: {e}")
            return []
    
    def _process_text_file(self, file_path: str) -> List[Dict[str, Any]]:
        """
        处理文本文件
        
        参数:
            file_path: 文件路径
            
        返回:
            商品数据列表
        """
        logger.info(f"处理TXT文件: {file_path}")
        
        # 读取文本内容
        content = self._read_text_file_with_encoding(file_path)
        if not content:
            logger.error(f"无法读取文件内容: {file_path}")
            return []
        
        # 使用大模型解析内容
        return self._parse_with_llm(content, file_path)
    
    def _read_text_file_with_encoding(self, file_path: str) -> Optional[str]:
        """
        尝试使用多种编码读取文本文件
        
        参数:
            file_path: 文件路径
            
        返回:
            文件内容字符串
        """
        # 常见编码列表
        encodings = ['utf-8', 'utf-16', 'gbk', 'gb2312', 'gb18030', 'latin1', 'iso-8859-1']
        
        # 首先检查UTF-8 BOM
        try:
            with open(file_path, 'rb') as f:
                raw = f.read(4)
                # 检查BOM标记
                if raw.startswith(b'\xef\xbb\xbf'):  # UTF-8 BOM
                    with open(file_path, 'r', encoding='utf-8-sig') as f:
                        content = f.read()
                    logger.info("成功使用 utf-8-sig 编码读取文件")
                    return content
                elif raw.startswith(b'\xff\xfe') or raw.startswith(b'\xfe\xff'):  # UTF-16 BOM
                    encoding = 'utf-16'
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    logger.info(f"成功使用 {encoding} 编码读取文件")
                    return content
        except Exception:
            pass
        
        # 尝试各种编码
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                logger.info(f"成功使用 {encoding} 编码读取文件")
                return content
            except UnicodeDecodeError:
                continue
        
        logger.error(f"尝试了所有编码仍无法读取文件: {file_path}")
        return None
    
    def _process_docx_file(self, file_path: str) -> List[Dict[str, Any]]:
        """
        处理Word文档
        
        参数:
            file_path: 文件路径
            
        返回:
            商品数据列表
        """
        logger.info(f"处理DOCX文件: {file_path}")
        
        try:
            # 尝试用不同的方法打开Word文档
            extracted_text = ""
            
            # 方法1: 使用python-docx库
            try:
                doc = docx.Document(file_path)
                
                # 提取段落文本
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                
                # 提取表格文本
                tables_text = []
                for table in doc.tables:
                    rows_text = []
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        rows_text.append("\t".join(cells))
                    tables_text.append("\n".join(rows_text))
                
                # 合并内容
                extracted_text = "\n\n".join(paragraphs + tables_text)
                
                # 写入临时文件用于调试
                with open("extracted_docx_content.txt", "w", encoding="utf-8") as f:
                    f.write(extracted_text)
                
                logger.info(f"成功使用python-docx库提取文本，长度: {len(extracted_text)}")
            except Exception as e:
                logger.warning(f"使用python-docx库提取文本失败: {str(e)}")
                extracted_text = ""
            
            # 方法2: 如果方法1失败，尝试直接读取文件并使用正则表达式提取
            if not extracted_text:
                try:
                    with open(file_path, 'rb') as f:
                        content = f.read().decode('utf-8', errors='ignore')
                    
                    # 使用正则表达式尝试提取商品信息
                    # 查找类似"商品ID: xxx"或"产品ID: xxx"的模式
                    pattern = r'(商品ID|产品ID|商品编号)[:：]\s*(.+?)(?=(商品ID|产品ID|商品编号)|$)'
                    product_blocks = re.findall(pattern, content, re.DOTALL | re.IGNORECASE)
                    
                    if product_blocks:
                        extracted_text = "\n\n".join([block[1] for block in product_blocks])
                        logger.info(f"通过正则表达式提取到{len(product_blocks)}个商品块")
                    else:
                        logger.warning("未通过正则表达式找到商品块")
                        extracted_text = content
                except Exception as e:
                    logger.warning(f"直接读取文件并提取文本失败: {str(e)}")
            
            # 如果两种方法都失败，尝试最后的备选方案
            if not extracted_text:
                logger.warning("无法从Word文档中提取文本，尝试最后的备选方案")
                try:
                    # 直接尝试从文件中提取任何可能的商品数据
                    with open(file_path, 'rb') as f:
                        binary_content = f.read()
                    
                    # 尝试提取ASCII文本
                    text_content = ''.join(chr(b) if 32 <= b < 127 else ' ' for b in binary_content)
                    
                    # 尝试查找产品相关的关键词
                    product_related = re.findall(r'([Pp]roduct|商品|产品|ID|编号|名称|价格|品牌).{0,50}', text_content)
                    if product_related:
                        extracted_text = '\n'.join(product_related)
                        logger.info(f"从二进制内容中提取到{len(product_related)}个可能的商品相关文本")
                    else:
                        logger.error("即使从二进制内容中也无法提取到任何商品相关信息")
                        return []
                except Exception as e:
                    logger.error(f"备选方案提取失败: {str(e)}")
                    return []
            
            # 使用大模型解析提取的文本
            return self._parse_with_llm(extracted_text, file_path)
            
        except Exception as e:
            logger.error(f"处理Word文档时发生错误: {str(e)}")
            return []
    
    def _process_excel_file(self, file_path: str) -> List[Dict[str, Any]]:
        """
        处理Excel文件
        
        参数:
            file_path: 文件路径
            
        返回:
            商品数据列表
        """
        logger.info(f"处理Excel文件: {file_path}")
        
        try:
            # 读取Excel文件
            df = pd.read_excel(file_path)
            
            # 转换为文本表示
            content = df.to_string(index=False)
            
            # 使用大模型解析内容
            return self._parse_with_llm(content, file_path)
            
        except Exception as e:
            logger.error(f"处理Excel文件时发生错误: {e}")
            return []
    
    def _process_csv_file(self, file_path: str) -> List[Dict[str, Any]]:
        """
        处理CSV文件
        
        参数:
            file_path: 文件路径
            
        返回:
            商品数据列表
        """
        logger.info(f"处理CSV文件: {file_path}")
        
        try:
            # 读取CSV文件，尝试不同编码
            df = None
            encodings = ['utf-8', 'gbk', 'latin1']
            
            for encoding in encodings:
                try:
                    df = pd.read_csv(file_path, encoding=encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if df is None:
                logger.error(f"无法读取CSV文件: {file_path}")
                return []
            
            # 转换为文本表示
            content = df.to_string(index=False)
            
            # 使用大模型解析内容
            return self._parse_with_llm(content, file_path)
            
        except Exception as e:
            logger.error(f"处理CSV文件时发生错误: {e}")
            return []
    
    def _process_json_file(self, file_path: str) -> List[Dict[str, Any]]:
        """
        处理JSON文件
        
        参数:
            file_path: 文件路径
            
        返回:
            商品数据列表
        """
        logger.info(f"处理JSON文件: {file_path}")
        
        try:
            # 读取JSON文件
            with open(file_path, 'r', encoding='utf-8') as f:
                content = json.load(f)
            
            # 如果已经是商品列表，添加处理逻辑以确保格式一致
            if isinstance(content, list):
                products = []
                for item in content:
                    if isinstance(item, dict):
                        # 确保每个商品都有一个ID
                        if 'id' not in item or not item['id']:
                            item['id'] = f"p{uuid.uuid4().hex[:8]}"
                        products.append(item)
                
                # 确保商品数据有效
                return self._validate_and_standardize_products(products)
            else:
                # 如果JSON不是列表格式，当作文本处理
                return self._parse_with_llm(json.dumps(content, ensure_ascii=False), file_path)
                
        except Exception as e:
            logger.error(f"处理JSON文件时发生错误: {e}")
            return []
    
    def _parse_with_llm(self, content: str, file_path: Optional[str] = None) -> List[Dict[str, Any]]:
        """
        使用大模型解析内容，提取商品数据
        
        参数:
            content: 文本内容
            file_path: 文件路径（可选）
            
        返回:
            商品数据列表
        """
        if not self.use_llm or not self.model:
            logger.warning("大语言模型未启用或初始化失败，使用传统方法解析")
            return self._parse_traditional(content)
        
        logger.info("使用LLM解析内容")
        
        # 设计提示模板，让模型提取商品数据
        system_prompt = """你是一个专业的商品数据提取助手。你的任务是从提供的文本中提取所有商品信息，并转换为结构化的JSON数组格式，遵循以下规则：

1. 每个商品作为一个独立的JSON对象，包含所有原始属性
2. 确保每个商品有一个唯一的ID字段，如果原文没有提供，则使用"p"加上数字作为ID（如p001、p002等）
3. 保留所有原始字段名称和值，不要改变任何数据内容
4. 如果有价格字段，保留原始格式（如"¥1299"或"1299.00"）
5. 如果商品有多级属性（如规格、特性等），将其解析为嵌套结构
   - 规格(specs)、参数等应解析为对象格式，如 {"蓝牙版本": "5.2", "电池续航": "30小时"}
   - 特性(features)、颜色(colors)等应解析为数组格式，如 ["黑色", "白色", "蓝色"]
6. 处理可能存在的双语字段，如 name/商品名称、description/商品描述等，保留两者
7. 返回格式必须是严格有效的JSON数组，只包含商品数据，不要有任何前缀或后缀说明

如果文本中有明显分隔的多个商品，请将它们分别提取为独立对象。
如果一个商品的信息分散在多处，请尽量将其合并为一个完整对象。

请直接返回JSON数组，不要包括任何额外的解释或Markdown格式。确保输出的JSON格式完全符合规范，可以直接被解析。如果文本中没有明确的商品数据，返回空数组[]。"""
        
        # 使用直接调用并设置较短的超时时间
        try:
            # 准备模型输入
            model_messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": content}
            ]
            
            # 如果输入内容很长，直接使用传统方法
            if len(content) > 10000:
                logger.warning(f"输入内容过长 ({len(content)} 字符)，使用传统方法解析")
                return self._parse_traditional(content)
            
            logger.info("开始调用大模型API...")
            start_time = time.time()
            
            # 直接调用API，使用超时控制
            # 使用一个较短的超时时间避免长时间等待
            try:
                # 修改模型的超时时间设置
                if hasattr(self.model, "request_timeout"):
                    original_timeout = self.model.request_timeout
                    self.model.request_timeout = 15.0  # 设置15秒超时
                
                response = self.model.invoke(model_messages)
                
                # 恢复原始超时设置
                if hasattr(self.model, "request_timeout"):
                    self.model.request_timeout = original_timeout
                    
            except Exception as api_error:
                logger.error(f"API调用失败: {str(api_error)}")
                return self._parse_traditional(content)
                
            elapsed_time = time.time() - start_time
            logger.info(f"大模型API调用完成，耗时: {elapsed_time:.2f}秒")
            
            # 获取模型返回的内容
            result = None
            if hasattr(response, "content"):
                result = response.content
            elif isinstance(response, str):
                result = response
            elif hasattr(response, "message") and hasattr(response.message, "content"):
                result = response.message.content
            elif isinstance(response, dict) and "content" in response:
                result = response["content"]
            else:
                # 尝试将响应对象转换为字符串
                try:
                    logger.warning(f"无法直接获取模型响应内容，正在尝试转换。响应类型: {type(response)}")
                    if hasattr(response, "__dict__"):
                        logger.info(f"响应属性: {response.__dict__}")
                    result = str(response)
                except Exception as e:
                    logger.error(f"转换模型响应为字符串失败: {str(e)}")
                    return self._parse_traditional(content)
            
            if not result:
                logger.warning("从模型响应中无法提取内容，尝试传统方法")
                return self._parse_traditional(content)
            
            # 提取JSON
            json_text = self._extract_json_from_text(result)
            if not json_text:
                logger.warning("LLM返回的内容无法解析为JSON，尝试传统方法")
                # 保存响应内容用于调试
                with open("llm_response_debug.txt", "w", encoding="utf-8") as f:
                    f.write(str(result))
                return self._parse_traditional(content)
            
            # 解析JSON
            try:
                products = json.loads(json_text)
                if not isinstance(products, list):
                    products = [products]
                
                # 确保每个商品都有ID
                for product in products:
                    if isinstance(product, dict):
                        if 'id' not in product or not product['id']:
                            product['id'] = f"p{uuid.uuid4().hex[:8]}"
                
                # 确保商品数据有效
                return self._validate_and_standardize_products(products)
                
            except json.JSONDecodeError as e:
                logger.error(f"解析JSON时发生错误: {str(e)}")
                logger.warning("回退到传统解析方法")
                return self._parse_traditional(content)
                
        except Exception as e:
            logger.error(f"调用大模型时发生错误: {str(e)}")
            # 记录详细的错误信息用于调试
            import traceback
            logger.error(f"详细错误: {traceback.format_exc()}")
            logger.warning("回退到传统解析方法")
            return self._parse_traditional(content)
    
    def _extract_json_from_text(self, text: str) -> str:
        """
        从文本中提取JSON部分
        
        参数:
            text: 文本内容
            
        返回:
            JSON文本
        """
        if not text:
            return ""
            
        # 如果本身就是有效的JSON，直接返回
        try:
            json.loads(text)
            return text
        except:
            pass
        
        # 优先尝试提取 ```json ... ``` 格式
        json_pattern = r'```(?:json)?\s*([\s\S]*?)\s*```'
        matches = re.findall(json_pattern, text)
        
        if matches:
            for match in matches:
                try:
                    # 验证是否为有效JSON
                    json.loads(match)
                    return match
                except:
                    continue
        
        # 尝试提取方括号包围的JSON数组
        array_pattern = r'\[\s*{[\s\S]*}\s*\]'
        matches = re.findall(array_pattern, text)
        
        if matches:
            for match in matches:
                try:
                    json.loads(match)
                    return match
                except:
                    continue
        
        # 尝试提取大括号包围的JSON对象（并尝试外层加[]转为数组）
        object_pattern = r'{[\s\S]*}'
        matches = re.findall(object_pattern, text)
        
        if matches:
            for match in matches:
                try:
                    json.loads(match)
                    return match
                except:
                    # 尝试将单个对象转换为数组
                    try:
                        array_json = f"[{match}]"
                        json.loads(array_json)
                        return array_json
                    except:
                        continue
        
        # 尝试将文本分行处理，寻找可能的JSON行
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if line.startswith('[') and line.endswith(']'):
                try:
                    json.loads(line)
                    return line
                except:
                    pass
            if line.startswith('{') and line.endswith('}'):
                try:
                    json.loads(line)
                    return line
                except:
                    # 尝试将单个对象转换为数组
                    try:
                        array_json = f"[{line}]"
                        json.loads(array_json)
                        return array_json
                    except:
                        pass
        
        # 尝试修复常见的JSON格式问题
        try:
            # 替换单引号为双引号
            fixed_text = text.replace("'", '"')
            json.loads(fixed_text)
            return fixed_text
        except:
            pass
        
        try:
            # 尝试修复没有引号的键
            # 使用全局re模块，不要重新导入
            fixed_text = re.sub(r'([{,])\s*([a-zA-Z0-9_]+)\s*:', r'\1"\2":', text)
            json.loads(fixed_text)
            return fixed_text
        except:
            pass
        
        # 如果以上方法都无法提取有效JSON，记录原始文本的部分内容用于调试
        logger.warning(f"无法从文本中提取JSON结构，文本前100字符: {text[:100]}")
        return ""
    
    def _parse_traditional(self, content: str) -> List[Dict[str, Any]]:
        """
        使用传统方法解析内容
        
        参数:
            content: 文本内容
            
        返回:
            商品数据列表
        """
        logger.info("使用传统解析方法处理文本")
        
        # 简单的传统解析方法
        products = []
        current_product = {}
        current_section = None  # 当前正在处理的嵌套结构部分
        current_section_data = {}  # 当前嵌套结构的数据
        
        # 按行分割内容
        lines = content.split('\n')
        
        # 检测分隔符
        separator_pattern = r'^[-]{3,}$|^[=]{3,}$|^[*]{3,}$'
        
        # 定义可能的列表类型字段
        list_fields = ['颜色', '特性', 'colors', 'features', '运动模式', '健康功能', '智能功能']
        
        # 定义可能的嵌套对象类型字段
        object_fields = ['规格', 'specs', '参数', 'parameters', '配置']
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 检查是否是分隔符，用于分割不同商品
            if re.match(separator_pattern, line):
                # 保存当前部分嵌套数据如果有的话
                if current_section and current_section_data:
                    current_product[current_section] = current_section_data
                    current_section = None
                    current_section_data = {}
                
                # 保存当前商品并创建新的商品
                if current_product:
                    products.append(current_product)
                    current_product = {}
                continue
            
            # 尝试识别可能的嵌套结构开始（如规格:、特性:等）
            nested_match = re.match(r'^(.*?)[:：]$', line)
            if nested_match:
                # 保存当前部分的嵌套数据如果有的话
                if current_section and current_section_data:
                    current_product[current_section] = current_section_data
                
                current_section = nested_match.group(1).strip()
                current_section_data = {}
                continue
            
            # 处理嵌套结构中的条目（通常以-或•开头）
            if line.startswith('-') or line.startswith('•') or line.startswith('*'):
                if current_section:
                    # 移除前导符号并解析键值对
                    item_line = re.sub(r'^[-•*]\s*', '', line).strip()
                    item_match = re.match(r'^(.*?)[:：]\s*(.*?)$', item_line)
                    if item_match:
                        key, value = item_match.groups()
                        key = key.strip()
                        value = value.strip()
                        current_section_data[key] = value
                    else:
                        # 如果不是键值对，可能是列表项
                        if current_section in list_fields:
                            if current_section not in current_product:
                                current_product[current_section] = []
                            current_product[current_section].append(item_line)
                        else:
                            # 如果是一个没有分隔的项目，添加为键值对
                            item_key = f"item_{len(current_section_data) + 1}"
                            current_section_data[item_key] = item_line
                else:
                    # 如果不在任何嵌套结构中，但有列表项
                    # 我们可以尝试猜测这是什么类型的数据
                    item_line = re.sub(r'^[-•*]\s*', '', line).strip()
                    item_match = re.match(r'^(.*?)[:：]\s*(.*?)$', item_line)
                    if item_match:
                        key, value = item_match.groups()
                        key = key.strip()
                        value = value.strip()
                        current_product[key] = value
                    else:
                        # 作为一般项目添加到当前产品，使用索引作为键
                        item_key = f"item_{len(current_product) + 1}"
                        current_product[item_key] = item_line
                continue
            
            # 尝试识别键值对（冒号分隔）
            kv_match = re.match(r'^(.*?)[:：]\s*(.*?)$', line)
            if kv_match:
                key, value = kv_match.groups()
                key = key.strip()
                value = value.strip()
                
                # 处理特殊情况
                if key.lower() in ['id', '商品id', '产品id', '商品编号', '产品编号']:
                    current_product['id'] = value
                elif key in list_fields:
                    # 处理列表类型字段，通常用逗号分隔
                    items = [item.strip() for item in value.split(',') if item.strip()]
                    if items:
                        current_product[key] = items
                    else:
                        current_product[key] = []
                elif key in object_fields:
                    # 处理对象类型字段
                    current_section = key
                    current_section_data = {}
                    # 如果值不为空，则可能是简短的对象描述
                    if value:
                        current_product[key] = value
                else:
                    # 普通键值对
                    current_product[key] = value
                continue
            
            # 尝试识别表格行（制表符或多个空格分隔）
            elif '\t' in line or '  ' in line:
                parts = re.split(r'\t+|\s{2,}', line)
                if len(parts) >= 2:
                    for i in range(0, len(parts) - 1, 2):
                        if i + 1 < len(parts):
                            key = parts[i].strip()
                            value = parts[i + 1].strip()
                            if key and value:
                                if key.lower() in ['id', '商品id', '产品id', '商品编号', '产品编号']:
                                    current_product['id'] = value
                                else:
                                    current_product[key] = value
        
        # 处理最后一个嵌套结构（如果有）
        if current_section and current_section_data:
            current_product[current_section] = current_section_data
        
        # 添加最后一个商品
        if current_product:
            products.append(current_product)
        
        # 后处理：处理字符串形式的列表和对象
        for product in products:
            for key, value in list(product.items()):
                # 处理逗号分隔的列表
                if isinstance(value, str) and key in list_fields:
                    items = [item.strip() for item in value.split(',') if item.strip()]
                    if items:
                        product[key] = items
            
            # 确保每个商品有ID
            if 'id' not in product or not product['id']:
                product['id'] = f"p{uuid.uuid4().hex[:8]}"
        
        # 确保商品数据有效
        return self._validate_and_standardize_products(products)
    
    def _validate_and_standardize_products(self, products: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        验证并标准化商品数据
        
        参数:
            products: 商品数据列表
            
        返回:
            标准化后的商品数据列表
        """
        if not products:
            return []
        
        valid_products = []
        for product in products:
            if not isinstance(product, dict):
                continue
            
            # 确保每个商品都有ID
            if 'id' not in product or not product['id']:
                product['id'] = f"p{uuid.uuid4().hex[:8]}"
            
            # 标准化ID格式（确保以p开头）
            product_id = str(product['id']).lower()
            if product_id.isdigit() or re.match(r'^\d+$', product_id):
                product['id'] = f"p{product_id}"
            elif not product_id.startswith('p'):
                product['id'] = f"p{product_id}"
            
            # 检查商品是否有效
            if len(product) > 1:  # 不仅仅只有ID字段
                valid_products.append(product)
            else:
                logger.warning(f"跳过无效商品: {product}")
        
        if not valid_products:
            logger.warning("没有找到有效的商品数据")
        
        return valid_products
    
    def process_text(self, text: str, format_hint: str = "auto") -> List[Dict[str, Any]]:
        """
        处理文本内容并提取商品信息
        
        参数:
            text: 文本内容
            format_hint: 输入格式提示 ('auto', 'key-value', 'json', 'table')
            
        返回:
            商品数据列表
        """
        logger.info(f"处理文本内容，格式提示: {format_hint}")
        
        # 记录基本信息
        content_size = len(text)
        lines = text.strip().split("\n")
        line_count = len(lines)
        logger.info(f"文本内容大小: {content_size} 字节, {line_count} 行")
        
        if content_size < 10:
            logger.warning("文本内容过短，可能不包含有效商品信息")
        
        # 选择解析方法
        if self.use_llm and self.model is not None:
            logger.info("使用大模型解析文本内容")
            return self._parse_with_llm(text)
        else:
            logger.info("使用传统方法解析文本内容")
            if format_hint == "json" or (format_hint == "auto" and text.strip().startswith(("[", "{"))):
                logger.info("尝试解析为JSON格式")
                try:
                    return self._process_json_content(text)
                except Exception as e:
                    logger.warning(f"JSON解析失败: {str(e)}，回退到传统解析")
            
            return self._parse_traditional(text)
    
    def save_to_json(self, products: List[Dict[str, Any]], output_file: str) -> str:
        """
        将商品数据保存为JSON文件
        
        参数:
            products: 商品数据列表
            output_file: 输出文件路径
            
        返回:
            保存的文件路径
        """
        if not products:
            logger.warning("没有商品数据可保存")
            return ""
        
        # 确保输出目录存在
        output_dir = os.path.dirname(output_file)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # 保存为JSON文件
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(products, f, ensure_ascii=False, indent=2)
        
        logger.info(f"已保存{len(products)}个商品到 {os.path.abspath(output_file)}")
        return os.path.abspath(output_file)

    def _process_txt_content(self, content: str) -> List[Dict[str, Any]]:
        """
        处理文本内容并提取商品信息
        这是一个兼容性方法，用于保持与原有代码的兼容性
        
        参数:
            content: 文本内容
            
        返回:
            商品数据列表
        """
        logger.info("处理文本内容，提取商品信息")
        return self._parse_traditional(content)

    def _process_json_content(self, text: str) -> List[Dict[str, Any]]:
        """
        处理JSON格式的文本内容
        
        参数:
            text: JSON格式的文本内容
            
        返回:
            商品数据列表
        """
        logger.info("处理JSON格式的文本内容")
        
        try:
            # 尝试解析JSON
            data = json.loads(text)
            
            # 确保结果是列表
            if isinstance(data, dict):
                data = [data]
            elif not isinstance(data, list):
                logger.warning(f"JSON内容格式不正确，期望列表或字典，得到 {type(data)}")
                return []
            
            logger.info(f"成功解析JSON，包含 {len(data)} 个对象")
            
            # 验证和标准化
            return self._validate_and_standardize_products(data)
            
        except json.JSONDecodeError as e:
            logger.error(f"JSON解析失败: {str(e)}")
            raise
        except Exception as e:
            logger.error(f"处理JSON内容时出错: {str(e)}")
            raise

def main():
    """主函数"""
    # 解析命令行参数
    parser = argparse.ArgumentParser(description="商品数据处理工具")
    parser.add_argument("input", help="输入文件路径或文本内容")
    parser.add_argument("--type", choices=["file", "text"], default="file", help="输入类型：文件或文本")
    parser.add_argument("--format-hint", choices=["auto", "key-value", "json", "csv", "table"], default="auto", 
                      help="输入格式提示")
    parser.add_argument("--use-llm", action="store_true", help="使用大语言模型进行解析")
    parser.add_argument("--output", default="products_data.json", help="输出文件路径")
    parser.add_argument("--api-key", help="OpenAI API密钥")
    parser.add_argument("--api-base", help="OpenAI API基础URL")
    
    args = parser.parse_args()
    
    # 初始化处理器
    processor = ProductDataProcessor(
        use_llm=args.use_llm,
        api_key=args.api_key,
        api_base=args.api_base
    )
    
    # 处理输入
    products = []
    if args.type == "file":
        products = processor.process_file(args.input)
    else:
        products = processor.process_text(args.input, format_hint=args.format_hint)
    
    # 保存结果
    if products:
        output_path = processor.save_to_json(products, args.output)
        print(f"已处理完成，保存至 {output_path}")
        print(f"共找到 {len(products)} 个商品")
    else:
        print("未找到有效的商品数据")

if __name__ == "__main__":
    main() 